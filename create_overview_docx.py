from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_overview_docx():
    doc = Document()
    
    # Title
    title = doc.add_heading('Document-Based Q&A Chatbot System - Technical Overview', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Project Architecture
    doc.add_heading('Project Architecture', 1)
    doc.add_paragraph('This is a document-based Q&A chatbot system with a modern web interface. The system allows users to upload documents (particularly PDFs), processes them, and then answers questions based on the content of these documents.')
    
    # System Components
    doc.add_heading('System Components', 2)
    components = [
        'Backend (Python/FastAPI)',
        'Frontend (Next.js/React)',
        'Vector Database (ChromaDB)',
        'Language Model (Gemma via Google API)'
    ]
    for component in components:
        doc.add_paragraph(component, style='List Bullet')
    
    # Technical Stack
    doc.add_heading('Technical Stack', 1)
    
    # Backend Technologies
    doc.add_heading('Backend Technologies', 2)
    backend_tech = [
        'FastAPI (v0.104.1) - Modern, fast web framework',
        'LangChain (v0.1.0) - Framework for LLM applications',
        'ChromaDB (v0.4.18) - Vector database for embeddings',
        'PyPDF2 (v3.0.1) - PDF processing',
        'Sentence Transformers - Text embeddings',
        'Google AI API - LLM access'
    ]
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    # Frontend Technologies
    doc.add_heading('Frontend Technologies', 2)
    frontend_tech = [
        'Next.js 13+ (App Router)',
        'TypeScript',
        'Tailwind CSS',
        'React',
        'Axios - HTTP client'
    ]
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    # Core Components
    doc.add_heading('Core Components', 1)
    
    # Document Processing Pipeline
    doc.add_heading('1. Document Processing Pipeline', 2)
    pipeline = [
        'File Upload → Text Extraction → Chunking → Embedding Creation → Vector Storage',
        'Supports PDF documents with size limit of 10MB',
        'Uses RecursiveCharacterTextSplitter for efficient text chunking',
        'Chunk size: 1000 characters with 200 character overlap'
    ]
    for item in pipeline:
        doc.add_paragraph(item, style='List Bullet')
    
    # Question Answering System
    doc.add_heading('2. Question Answering System', 2)
    qa_system = [
        'Question → Embedding Creation → Context Retrieval → LLM Response',
        'Uses Gemma 3 (27B parameter model) via Google AI API',
        'Temperature setting: 0.1 for focused responses',
        'Implements RetrievalQA chain for context-aware responses',
        'Training data Q&A pairs for common HR/IT queries'
    ]
    for item in qa_system:
        doc.add_paragraph(item, style='List Bullet')
    
    # Vector Database
    doc.add_heading('3. Vector Database (ChromaDB)', 2)
    vector_db = [
        'Persistent storage for document embeddings',
        'Efficient similarity search',
        'Document chunking and indexing',
        'Separate collections for user documents and training data'
    ]
    for item in vector_db:
        doc.add_paragraph(item, style='List Bullet')
    
    # Embeddings System
    doc.add_heading('4. Embeddings System', 2)
    embeddings = [
        'Model: all-MiniLM-L6-v2',
        'Normalized embeddings for better similarity matching',
        'LRU caching for performance optimization',
        'Batch processing capabilities'
    ]
    for item in embeddings:
        doc.add_paragraph(item, style='List Bullet')
    
    # API Endpoints
    doc.add_heading('API Endpoints', 1)
    doc.add_heading('Backend Endpoints', 2)
    
    endpoints = [
        ('POST /upload', [
            'Handles document uploads',
            'Processes PDF files',
            'Creates and stores embeddings'
        ]),
        ('POST /ask', [
            'Processes user questions',
            'Retrieves relevant context',
            'Generates answers using LLM',
            'Combines training data with user-uploaded documents'
        ]),
        ('POST /clear', [
            'Clears the document collection',
            'Maintains training data collection'
        ]),
        ('GET /', [
            'Health check endpoint',
            'System status verification'
        ])
    ]
    
    for endpoint, details in endpoints:
        p = doc.add_paragraph()
        p.add_run(endpoint).bold = True
        for detail in details:
            doc.add_paragraph(detail, style='List Bullet')
    
    # Performance Optimizations
    doc.add_heading('Performance Optimizations', 1)
    
    optimizations = [
        ('Caching', [
            'LRU cache for embeddings',
            'Document chunk caching',
            'Response caching'
        ]),
        ('Processing', [
            'Batch processing for documents',
            'Efficient text splitting',
            'Optimized ChromaDB settings'
        ]),
        ('Error Handling', [
            'File size validation',
            'PDF processing errors',
            'API error responses with detailed logging',
            'Frontend error display'
        ])
    ]
    
    for category, items in optimizations:
        doc.add_heading(category, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # Security Features
    doc.add_heading('Security Features', 1)
    
    security = [
        ('File Validation', [
            'Size limits (10MB)',
            'File type checking (.txt, .pdf, .doc, .docx)',
            'Secure file processing'
        ]),
        ('API Security', [
            'CORS middleware',
            'Error handling',
            'Input validation'
        ])
    ]
    
    for category, items in security:
        doc.add_heading(category, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # Technical Implementation Details
    doc.add_heading('Technical Implementation Details', 1)
    
    # Document Processing
    doc.add_heading('Document Processing', 2)
    code = doc.add_paragraph()
    code.add_run('''# Key parameters
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB''').font.name = 'Courier New'
    
    # Vector Database Configuration
    doc.add_heading('Vector Database Configuration', 2)
    code = doc.add_paragraph()
    code.add_run('''# ChromaDB settings
Settings(
    anonymized_telemetry=False,
    allow_reset=True,
    is_persistent=True
)''').font.name = 'Courier New'
    
    # LLM Configuration
    doc.add_heading('LLM Configuration', 2)
    code = doc.add_paragraph()
    code.add_run('''# Google AI API settings
llm = ChatOpenAI(
    model="gemma3-27b",
    temperature=0.1,
    openai_api_key=os.getenv("GOOGLE_API_KEY"),
    openai_api_base="https://litellm.dev.ai-cloud.me/v1"
)''').font.name = 'Courier New'
    
    # Training Data
    doc.add_heading('Training Data', 2)
    training_data = [
        'HR FAQs (time off, working hours, benefits)',
        'IT FAQs (password reset, hardware issues, software requests)',
        'Basic conversation patterns'
    ]
    for item in training_data:
        doc.add_paragraph(item, style='List Bullet')
    
    # Frontend Features
    doc.add_heading('Frontend Features', 1)
    
    frontend_features = [
        ('Modern UI', [
            'Responsive design with Tailwind CSS',
            'Dark theme with gradient accents',
            'Real-time server status indicator'
        ]),
        ('User Experience', [
            'File upload progress indicators',
            'Loading states and animations',
            'Clear error messaging',
            'Chat history with distinct user/bot messages'
        ]),
        ('Chat Management', [
            'New chat functionality',
            'File upload interface',
            'Message scrolling with auto-focus'
        ])
    ]
    
    for category, items in frontend_features:
        doc.add_heading(category, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # Common Technical Questions and Answers
    doc.add_heading('Common Technical Questions and Answers', 1)
    qa_pairs = [
        ('Why ChromaDB?', 'Offers persistent storage, LangChain integration, and efficient similarity search.'),
        ('Document chunking strategy?', 'Uses RecursiveCharacterTextSplitter with 1000/200 size/overlap for context preservation.'),
        ('Large document handling?', 'Implements 10MB limit, chunked processing, and efficient text splitting.'),
        ('Answer relevance?', 'Combines semantic search and RetrievalQA chain for context-aware responses.'),
        ('Caching strategy?', 'LRU caching for embeddings to improve performance and reduce computation.'),
        ('PDF processing?', 'PyPDF2 for extraction with async processing and error handling.'),
        ('Security measures?', 'File size limits, type validation, CORS, and proper error handling.'),
        ('Frontend-Backend communication?', 'RESTful API endpoints with Axios, timeout handling, and server status checks.')
    ]
    
    for question, answer in qa_pairs:
        p = doc.add_paragraph()
        p.add_run(f'Q: {question}').bold = True
        doc.add_paragraph(f'A: {answer}')
    
    # Development Setup
    doc.add_heading('Development Setup', 1)
    
    # Backend Setup
    doc.add_heading('1. Backend Setup', 2)
    code = doc.add_paragraph()
    code.add_run('''pip install -r requirements.txt
uvicorn main:app --reload''').font.name = 'Courier New'
    
    # Frontend Setup
    doc.add_heading('2. Frontend Setup', 2)
    code = doc.add_paragraph()
    code.add_run('''npm install
npm run dev''').font.name = 'Courier New'
    
    # Best Practices
    doc.add_heading('Best Practices Implemented', 1)
    
    practices = [
        ('Code Organization', [
            'Modular architecture',
            'Clear separation of concerns',
            'Proper error handling and logging'
        ]),
        ('Performance', [
            'Efficient document processing',
            'Optimized embeddings',
            'Caching strategies'
        ]),
        ('User Experience', [
            'Responsive design',
            'Real-time feedback',
            'Error handling and server status indicators'
        ]),
        ('Security', [
            'Input validation',
            'File processing safety',
            'API security measures'
        ])
    ]
    
    for category, items in practices:
        doc.add_heading(category, 2)
        for item in items:
            doc.add_paragraph(item, style='List Bullet')
    
    # Save the document
    doc.save('data/project_overview.docx')

if __name__ == '__main__':
    create_overview_docx() 